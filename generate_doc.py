# -*- coding: utf-8 -*-
"""Generate the full SmartDine (ToDining) project explanation Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Theme colors ----------
PRIMARY = RGBColor(0xC0, 0x39, 0x2B)   # restaurant red
DARK    = RGBColor(0x1F, 0x2A, 0x37)   # slate
ACCENT  = RGBColor(0xE6, 0x7E, 0x22)   # orange
GREY    = RGBColor(0x55, 0x5F, 0x6B)
LIGHT   = RGBColor(0xFA, 0xF3, 0xEF)

doc = Document()

# ---------- Base styles ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=10, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = PRIMARY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'E67E22')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = DARK
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT
    return p


def body(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(label, text, fill='FAF3EF', barcolor=PRIMARY):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + '  ')
    r.bold = True
    r.font.color.rgb = barcolor
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    doc.add_paragraph()
    return tbl


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], 'C0392B')
        set_cell_text(hdr[i], htext, bold=True, white=True, size=10)
    for r_i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
            if r_i % 2 == 1:
                shade(cells[i], 'F7F1ED')
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SmartDine')
r.bold = True
r.font.size = Pt(54)
r.font.color.rgb = PRIMARY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('( ToDining )')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run('Smart Restaurant Management SaaS Platform')
r.font.size = Pt(15)
r.font.color.rgb = GREY

doc.add_paragraph()
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run('Complete Project Explanation  ' + chr(0x2022) + '  Step-by-Step Guide')
r.font.size = Pt(13)
r.bold = True
r.font.color.rgb = DARK

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run('QR Ordering  ' + chr(0x2022) + '  Kitchen & Waiter Boards  ' + chr(0x2022) +
                 '  Reservations  ' + chr(0x2022) + '  Billing  ' + chr(0x2022) +
                 '  Inventory  ' + chr(0x2022) + '  Analytics  ' + chr(0x2022) + '  Multi-Restaurant')
r.font.size = Pt(10.5)
r.italic = True
r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Document Version 1.0\nPrepared: June 2026')
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
h1('Table of Contents')
toc_items = [
    '1.  Project Overview',
    '2.  Technology Stack',
    '3.  System Architecture (How It All Fits Together)',
    '4.  User Roles & Who Can See What',
    '5.  The Four Surfaces of SmartDine',
    '6.  Customer Journey (QR Ordering) - Step by Step',
    '7.  Staff Dashboards - Explained One by One',
    '       7.1  Kitchen Dashboard',
    '       7.2  Waiter Dashboard',
    '8.  Admin Panel Dashboards - Explained One by One',
    '       8.1  Main Admin Dashboard',
    '       8.2  Analytics Dashboard',
    '       8.3  Orders',
    '       8.4  Tables & QR',
    '       8.5  Menu Management',
    '       8.6  Categories',
    '       8.7  Reservations',
    '       8.8  Inventory',
    '       8.9  Billing',
    '       8.10 Feedback',
    '       8.11 Notifications',
    '       8.12 Staff',
    '       8.13 Restaurants (Super-Admin)',
    '9.  Database & Backend (Supabase)',
    '10. Complete Order Lifecycle (End to End)',
    '11. How to Run the Project',
    '12. Glossary',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    indent = item.startswith('     ')
    r = p.add_run(item.strip())
    r.font.size = Pt(11)
    if not indent and item[0].isdigit():
        r.bold = True
        r.font.color.rgb = DARK
    else:
        r.font.color.rgb = GREY
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)

doc.add_page_break()

# ============================================================
# 1. PROJECT OVERVIEW
# ============================================================
h1('1.  Project Overview')
body('SmartDine (internally named "ToDining") is a complete, mobile-first software platform '
     'that runs the day-to-day operations of a restaurant. It replaces paper menus, handwritten '
     'order tickets, manual billing, and disconnected reservation books with one connected '
     'digital system that every part of the restaurant shares in real time.')

body('It is built as a multi-tenant SaaS (Software-as-a-Service) platform. This means a single '
     'installation can run many different restaurants at the same time, with each restaurant\'s '
     'data kept completely separate and private from the others.')

h2('What problem does it solve?')
body('In a normal restaurant, a customer waits for a waiter to take an order, the waiter walks '
     'the ticket to the kitchen, the kitchen cooks, someone tracks stock by hand, and the bill is '
     'calculated manually. Information lives in different places and gets lost. SmartDine connects '
     'all of these steps so that the moment a customer places an order on their phone, the kitchen, '
     'the waiters, and the manager all see it instantly.')

h2('The 16 core features')
features = [
    ('QR Table Ordering', 'Customers scan a QR code on the table and order directly from their phone - no app, no login.'),
    ('Digital Menu', 'Menu with categories, photos, prices, tags, and live availability toggling.'),
    ('Order Lifecycle', 'Every order moves through clear stages: pending -> preparing -> ready -> served -> completed.'),
    ('Kitchen Board', 'A live Kanban screen where cooks see and advance every order ticket.'),
    ('Waiter Board', 'A live screen showing which dishes are ready to serve and which tables need help.'),
    ('Waiter Call System', 'Guests can request a waiter, water, the bill, or assistance with one tap.'),
    ('Reservations', 'A public booking form plus a management screen to confirm or reject bookings.'),
    ('Live Table Status', 'Each table is shown as available, reserved, or occupied in real time.'),
    ('Billing', 'Automatic bills with tax and service charge, printable and exportable as PDF.'),
    ('Feedback', 'Customers rate food, service, and overall experience after their meal.'),
    ('Inventory', 'Stock tracking that auto-deducts ingredients when dishes are ordered, with low-stock alerts.'),
    ('AI Upselling', 'A suggestion engine that recommends add-ons (e.g. "add a drink with that").'),
    ('WhatsApp Notifications', 'A service that prepares and logs WhatsApp messages to customers.'),
    ('Analytics Dashboard', 'Revenue, peak hours, top-selling dishes, and rating trends in charts.'),
    ('Multi-Restaurant SaaS', 'One system runs many restaurants with full data isolation.'),
    ('Real-time Sync', 'Every screen updates instantly across customer, kitchen, waiter, and admin.'),
]
make_table(['Feature', 'What it does'], features, widths=[1.9, 4.6])

callout('IN SHORT:', 'SmartDine is a restaurant\'s complete digital nervous system - from the '
        'customer scanning a QR code at the table, all the way to the owner viewing revenue '
        'analytics, everything is connected and updates live.')

doc.add_page_break()

# ============================================================
# 2. TECH STACK
# ============================================================
h1('2.  Technology Stack')
body('The application is built entirely with modern web technologies, so it runs in any browser '
     'on any phone, tablet, or computer with nothing to install.')

h2('Frontend (what the user sees and interacts with)')
make_table(['Technology', 'Role in the project'], [
    ('React 18 + TypeScript', 'The core framework that builds the user interface; TypeScript adds safety to the code.'),
    ('Vite 6', 'The build tool and development server - makes the app fast to develop and load.'),
    ('React Router v6', 'Handles navigation between pages and the different dashboards.'),
    ('Tailwind CSS 3', 'The styling system that gives the app its clean, consistent look.'),
    ('Zustand + React Context', 'Manages shared application state (current user, active restaurant, cart).'),
    ('React Hook Form + Zod', 'Builds and validates all the forms (login, reservations, menu editing).'),
    ('Recharts', 'Draws the charts on the Analytics dashboard.'),
    ('jsPDF', 'Generates printable PDF bills.'),
    ('qrcode', 'Generates the QR codes for each table.'),
    ('lucide-react / sonner', 'Icon set and on-screen toast notifications (the live alerts).'),
], widths=[2.1, 4.4])

h2('Backend (where data is stored and secured)')
make_table(['Technology', 'Role in the project'], [
    ('Supabase', 'The cloud backend - PostgreSQL database, user authentication, and real-time updates.'),
    ('PostgreSQL', 'The relational database that stores every restaurant, order, menu item, etc.'),
    ('Row Level Security (RLS)', 'Database rules that guarantee one restaurant can never see another\'s data.'),
    ('Vercel', 'The hosting platform the finished app is deployed to.'),
], widths=[2.1, 4.4])

callout('SMART DESIGN:', 'The app has a built-in fallback. If Supabase is not configured, it '
        'automatically runs on local browser storage with demo data - so development and demos '
        'never get blocked. Going live is simply a matter of adding the Supabase keys; no screens '
        'need to be rewritten.')

doc.add_page_break()

# ============================================================
# 3. ARCHITECTURE
# ============================================================
h1('3.  System Architecture (How It All Fits Together)')
body('SmartDine is organised in clean layers. Understanding these layers makes every dashboard '
     'easier to follow, because they all share the same foundation.')

h2('The layered structure')
numbered('The screens (pages) - thin components that the user actually sees. They never talk to '
         'the database directly.', bold_lead='Presentation layer: ')
numbered('A single set of "services" (menu service, order service, billing service, etc.). Every '
         'screen reads and writes through these services only.', bold_lead='Data service layer: ')
numbered('Either the live Supabase database OR local browser storage with demo data - chosen '
         'automatically by configuration.', bold_lead='Storage layer: ')
numbered('An event bus + live-query hook that pushes changes to every open screen instantly, '
         'simulating (or using) Supabase Realtime.', bold_lead='Real-time layer: ')

h2('Why this matters')
body('Because every screen goes through the same service layer, the same piece of information '
     '(for example, a new order) appears at the same moment on the customer\'s phone, the kitchen '
     'screen, the waiter screen, and the admin dashboard. There is one single source of truth, and '
     'everyone is always looking at the same up-to-date picture.')

h2('Main project folders')
make_table(['Folder', 'What lives there'], [
    ('src/app', 'The app shell: providers, the router (all routes), and the entry point.'),
    ('src/components/layout', 'The frames around screens: admin sidebar, staff layout, customer layout, role guard.'),
    ('src/components/ui', 'Reusable building blocks: buttons, cards, inputs, modals, badges, KPI cards.'),
    ('src/features', 'Feature modules: menu, cart, orders, kitchen, waiter, reservations, billing, etc.'),
    ('src/data/services', 'The data service layer - one service per domain.'),
    ('src/data/mock', 'Seed data and the in-memory store used for the demo/fallback mode.'),
    ('src/data/supabase', 'The live Supabase client and data mappers.'),
    ('src/context', 'Shared state: Auth (who is logged in), Tenant (active restaurant), Cart.'),
    ('src/pages', 'The thin route components grouped into admin / staff / customer.'),
    ('supabase', 'The database schema, security policies, and setup scripts.'),
], widths=[2.0, 4.5])

doc.add_page_break()

# ============================================================
# 4. ROLES
# ============================================================
h1('4.  User Roles & Who Can See What')
body('SmartDine has four staff roles plus the anonymous customer. Each role automatically lands '
     'on the screen most relevant to their job and can only access what they are allowed to.')

make_table(['Role', 'Lands on', 'Can access', 'Special powers'], [
    ('Owner', 'Analytics', 'Everything', 'Sees Analytics + manages all restaurants (super-admin).'),
    ('Manager', 'Orders', 'Admin + Kitchen + Waiter', 'Runs the restaurant day-to-day; no owner-only screens.'),
    ('Waiter', 'Waiter board', 'Waiter board only', 'Serves food, answers service calls.'),
    ('Kitchen', 'Kitchen board', 'Kitchen board only', 'Cooks and advances order tickets.'),
    ('Customer', 'Menu (via QR)', 'Order & track only', 'No login - identified by the table QR code.'),
], widths=[1.0, 1.2, 1.8, 2.5])

callout('SECURITY NOTE:', 'Access is enforced in two places. On the screen, a "Role Guard" blocks '
        'pages a user is not allowed to open. In the database, Row Level Security policies make it '
        'physically impossible for a user to read or change another restaurant\'s data, even if '
        'they tried to bypass the screen.')

h2('How login works')
body('Staff sign in with their email. In the demo build there are one-tap buttons for each of the '
     'four roles (Owner, Manager, Waiter, Kitchen), so you can instantly experience the app from '
     'each perspective. In production, each staff record is linked to a real Supabase Auth account '
     'with a password. Customers never log in - scanning the table\'s QR code is what identifies '
     'them and their table.')

doc.add_page_break()

# ============================================================
# 5. FOUR SURFACES
# ============================================================
h1('5.  The Four Surfaces of SmartDine')
body('The whole product is divided into four "surfaces" - distinct areas of the app, each designed '
     'for a different type of user. Every dashboard in this document belongs to one of these four.')

h2('1) Customer surface (public, no login)')
body('Mobile-first screens reached by scanning a table QR code. This is where guests browse the '
     'menu, order, track their food, pay, and leave feedback.')

h2('2) Staff surface (role-guarded boards)')
body('Two fast, live "boards" built for people on their feet: the Kitchen board for cooks and the '
     'Waiter board for servers.')

h2('3) Admin Panel (managers and owners)')
body('The full management back office - thirteen screens covering the dashboard, analytics, '
     'orders, tables, menu, categories, reservations, inventory, billing, feedback, notifications, '
     'staff, and restaurants. This is the heart of the system for running the business.')

h2('4) Public / Auth pages')
body('The marketing landing page, the staff login page, and a "not found" page.')

h2('Complete route map')
make_table(['Surface', 'Route', 'Screen'], [
    ('Public', '/', 'Landing page'),
    ('Public', '/login', 'Staff login'),
    ('Customer', '/r/:slug/t/:tableId', 'Menu page'),
    ('Customer', '/r/:slug/t/:tableId/cart', 'Cart page'),
    ('Customer', '/r/:slug/t/:tableId/order/:id', 'Order tracking page'),
    ('Customer', '/reserve/:slug', 'Public reservation form'),
    ('Staff', '/kitchen', 'Kitchen dashboard'),
    ('Staff', '/waiter', 'Waiter dashboard'),
    ('Admin', '/admin', 'Main admin dashboard'),
    ('Admin', '/admin/analytics', 'Analytics dashboard (owner)'),
    ('Admin', '/admin/orders', 'Orders'),
    ('Admin', '/admin/tables', 'Tables & QR'),
    ('Admin', '/admin/menu', 'Menu management'),
    ('Admin', '/admin/categories', 'Categories'),
    ('Admin', '/admin/reservations', 'Reservations'),
    ('Admin', '/admin/inventory', 'Inventory'),
    ('Admin', '/admin/billing', 'Billing'),
    ('Admin', '/admin/feedback', 'Feedback'),
    ('Admin', '/admin/notifications', 'Notifications'),
    ('Admin', '/admin/staff', 'Staff'),
    ('Admin', '/admin/restaurants', 'Restaurants (super-admin, owner)'),
], widths=[1.0, 2.6, 2.9])

doc.add_page_break()

# ============================================================
# 6. CUSTOMER JOURNEY
# ============================================================
h1('6.  Customer Journey (QR Ordering) - Step by Step')
body('This is the experience a guest has at the table. It needs no app and no account.')

numbered('The guest sits at a table and scans the QR code printed on it. This opens the menu page '
         'in their phone browser at a unique link that already knows which restaurant and which '
         'table they are at.', bold_lead='Scan the QR code: ')
numbered('The app confirms the table link is valid and shows the table number at the top. A "Help" '
         'button is always available.', bold_lead='Table is detected: ')
numbered('The guest scrolls the digital menu, organised into categories with photos, prices, and '
         'tags. Tapping a dish adds it to the cart. The upsell engine may suggest a matching add-on.',
         bold_lead='Browse the menu: ')
numbered('A cart bar shows the running total. On the cart page the guest reviews items and places '
         'the order with one tap.', bold_lead='Place the order: ')
numbered('Instantly, the order appears on the kitchen and waiter screens (see the lifecycle in '
         'section 10). The guest is taken to a live tracking page.', bold_lead='Order goes live: ')
numbered('The tracking page shows a live timeline - pending, preparing, ready, served - updating '
         'on its own as the kitchen advances the order.', bold_lead='Track the food: ')
numbered('At any time the guest can tap Help to request a waiter, water, the bill, or assistance. '
         'That request pops up on the waiter board.', bold_lead='Call for service: ')
numbered('When the meal is done, the bill (with tax and service charge) is shown, and the guest '
         'rates the food, service, and overall experience.', bold_lead='Pay & give feedback: ')

callout('KEY POINT:', 'The customer never sees a "dashboard" in the management sense - their '
        'experience is deliberately simple. But every tap they make feeds the staff and admin '
        'dashboards explained next.')

doc.add_page_break()

# ============================================================
# 7. STAFF DASHBOARDS
# ============================================================
h1('7.  Staff Dashboards - Explained One by One')
body('The two staff boards are designed to be glanced at from across a busy kitchen or floor. '
     'They update live and use colour and toasts (pop-up alerts) to grab attention.')

# 7.1 Kitchen
h2('7.1  Kitchen Dashboard')
body('Route: /kitchen  |  Who can use it: Kitchen staff, Managers, Owners')
h3('Purpose')
body('Gives cooks a single live board of every order ticket so nothing is missed and food is '
     'prepared in the right order.')
h3('What is on the screen')
bullet('A Kanban board with three columns: New, Preparing, and Ready.', bold_lead='Three columns: ')
bullet('Each card shows the table number, the items ordered, and quantities.', bold_lead='Order cards: ')
bullet('When a new order arrives, a toast alert pops up and the ticket appears in the New column.', bold_lead='Live alerts: ')
bullet('Buttons on each card move it from New -> Preparing -> Ready.', bold_lead='Advance buttons: ')
h3('How a cook uses it - step by step')
numbered('A new ticket appears in the New column with a sound/toast alert.')
numbered('The cook taps "Start" to move it to Preparing and begins cooking.')
numbered('When the dish is done, the cook taps "Ready". The ticket moves to the Ready column.')
numbered('Instantly, the waiter board lights up showing that this table\'s food is ready to serve.')

# 7.2 Waiter
h2('7.2  Waiter Dashboard')
body('Route: /waiter  |  Who can use it: Waiters, Managers, Owners')
h3('Purpose')
body('Tells servers exactly what to do next: which dishes to carry out and which tables are '
     'calling for help.')
h3('What is on the screen')
bullet('KPI cards at the top: active tables, dishes ready to serve, and open service calls.', bold_lead='Live counters: ')
bullet('A list of guest requests (waiter, water, bill, assistance) as they come in.', bold_lead='Service calls: ')
bullet('A "Ready to serve" section listing food the kitchen has finished.', bold_lead='Ready to serve: ')
bullet('An "In the kitchen" section showing what is still being prepared.', bold_lead='In progress: ')
bullet('Toast alerts when a new service call or ready dish appears.', bold_lead='Live alerts: ')
h3('How a waiter uses it - step by step')
numbered('A dish reaches the Ready column in the kitchen; it appears under "Ready to serve" here.')
numbered('The waiter carries the food to the correct table (the card shows the table number).')
numbered('The waiter marks it served; the order advances and the customer\'s tracking page updates.')
numbered('If a guest taps Help, a service call appears at the top; the waiter handles it and clears it.')

callout('WHY TWO BOARDS:', 'Splitting kitchen and waiter views keeps each role focused. Cooks only '
        'care about what to make; waiters only care about what to carry and which table needs them. '
        'Both are fed by the same live order data.')

doc.add_page_break()

# ============================================================
# 8. ADMIN DASHBOARDS
# ============================================================
h1('8.  Admin Panel Dashboards - Explained One by One')
body('The Admin Panel is the management back office. It is wrapped in a layout with a sidebar '
     '(a mobile drawer on phones) and a restaurant switcher at the top. Below, every one of the '
     'thirteen admin screens is explained in turn.')

admin_screens = [
    ('8.1  Main Admin Dashboard', '/admin', 'Managers, Owners',
     'The landing screen for managers - a quick health check of the whole restaurant at a glance.',
     [('KPI cards', 'Four headline numbers: total menu items, total orders, number of tables, and today\'s revenue.'),
      ('Recent orders', 'A live list of the most recent orders with their current status.'),
      ('Quick actions', 'Shortcut links to jump straight to menu, tables, orders, and other screens.')],
     ['Open the panel and immediately see how the restaurant is doing today.',
      'Glance at the KPI cards for the big picture (revenue, order count).',
      'Scan recent orders to spot anything that needs attention.',
      'Use a quick-action link to jump to the screen you need.']),

    ('8.2  Analytics Dashboard', '/admin/analytics', 'Owners only',
     'The owner\'s decision-making screen - turns raw operations into business insight with charts.',
     [('Revenue KPIs', 'Today, this week, this month, and average order value.'),
      ('7-day revenue chart', 'An area chart showing the revenue trend across the last week.'),
      ('Peak hours chart', 'A bar chart revealing the busiest times of day.'),
      ('Top foods', 'A ranked list of the best-selling dishes.'),
      ('Reservations & ratings', 'A pie chart of reservation outcomes plus average ratings.')],
     ['The owner signs in and lands here automatically.',
      'Check revenue KPIs to see how the business is trending.',
      'Read the peak-hours chart to plan staffing.',
      'Use top-foods to decide menu promotions and pricing.']),

    ('8.3  Orders', '/admin/orders', 'Managers, Owners',
     'A complete, filterable record of every order with full status control.',
     [('All orders list', 'Every order across the restaurant, newest first.'),
      ('Status management', 'Change or correct an order\'s status when needed.'),
      ('Order details', 'See items, quantities, totals, and the table for each order.')],
     ['Open Orders to see the full history and live queue.',
      'Filter or scan to find a specific order.',
      'Adjust an order\'s status if something needs correcting.',
      'Use it as the source of truth when a customer has a question.']),

    ('8.4  Tables & QR', '/admin/tables', 'Managers, Owners',
     'Manages the physical tables and generates the QR codes guests scan.',
     [('Live table status', 'Each table shown as available, reserved, or occupied in real time.'),
      ('QR generation', 'Create and view a printable QR code for every table.'),
      ('Table setup', 'Add tables and set how many seats each has.')],
     ['Add each physical table with its seat count.',
      'Generate the QR code for every table.',
      'Print and place the QR codes on the tables.',
      'Watch table status change live as guests are seated and served.']),

    ('8.5  Menu Management', '/admin/menu', 'Managers, Owners',
     'The full create/edit/delete control over every dish on the menu.',
     [('Menu item CRUD', 'Add, edit, and remove dishes.'),
      ('Pricing & details', 'Set price, description, image, and tags.'),
      ('Availability toggle', 'Mark a dish as available or sold out instantly.')],
     ['Add a new dish with its name, price, photo, and category.',
      'Edit details or pricing whenever they change.',
      'Toggle a dish off when it sells out - it disappears from the customer menu live.',
      'Remove dishes that are retired from the menu.']),

    ('8.6  Categories', '/admin/categories', 'Managers, Owners',
     'Organises the menu into sections (starters, mains, drinks, desserts...).',
     [('Category list', 'All menu sections in their display order.'),
      ('Add / edit / sort', 'Create sections and control the order they appear in.')],
     ['Create the categories the menu should be grouped into.',
      'Set their order so the menu reads logically.',
      'Assign dishes to categories from the Menu Management screen.']),

    ('8.7  Reservations', '/admin/reservations', 'Managers, Owners',
     'Receives and manages table bookings made through the public reservation form.',
     [('Booking list', 'All reservations with name, contact, date, time, and guest count.'),
      ('Confirm / reject', 'Approve or decline each booking.'),
      ('Status tracking', 'See which bookings are pending, confirmed, or cancelled.')],
     ['A guest submits the public reservation form.',
      'The booking appears here as "pending".',
      'The manager confirms or rejects it.',
      'The table status updates to "reserved" for the chosen time.']),

    ('8.8  Inventory', '/admin/inventory', 'Managers, Owners',
     'Tracks raw stock and automatically deducts ingredients as dishes are ordered.',
     [('Stock items', 'All ingredients with current quantity and unit.'),
      ('Low-stock alerts', 'Highlights items that have dropped below their threshold.'),
      ('Auto-deduct via recipes', 'When a dish is ordered, its recipe ingredients are subtracted automatically.')],
     ['Add each ingredient with its starting quantity and low-stock threshold.',
      'Link ingredients to dishes (recipes) so deductions are automatic.',
      'As orders come in, stock drops on its own.',
      'Restock when an item triggers a low-stock alert.']),

    ('8.9  Billing', '/admin/billing', 'Managers, Owners',
     'Holds the history of every bill and lets staff export them as PDF.',
     [('Bill history', 'Every completed bill with its full breakdown.'),
      ('Tax & service charge', 'Each bill shows subtotal, tax, service charge, and grand total.'),
      ('PDF export', 'Download or print any bill as a professional PDF.')],
     ['When an order completes, a bill is generated automatically.',
      'The bill includes subtotal, tax, and service charge.',
      'Find any past bill in the history.',
      'Export or reprint it as a PDF when needed.']),

    ('8.10  Feedback', '/admin/feedback', 'Managers, Owners',
     'Collects the ratings and comments customers leave after their meal.',
     [('Ratings list', 'Food, service, and experience scores (1-5) per visit.'),
      ('Comments', 'Free-text feedback from guests.'),
      ('Trends', 'A view of how satisfaction is trending over time.')],
     ['A customer rates their visit on the tracking page.',
      'The rating and comment appear here.',
      'Management reviews feedback to spot problems and praise.',
      'Recurring issues guide improvements to food and service.']),

    ('8.11  Notifications', '/admin/notifications', 'Managers, Owners',
     'A log and preview of the WhatsApp messages the system sends to customers.',
     [('Message log', 'Every notification with recipient, type, and status.'),
      ('Previews', 'See exactly what the customer message looks like.'),
      ('Channel ready', 'Built to connect to WhatsApp providers (Meta / Twilio).')],
     ['The system prepares a WhatsApp message (e.g. order ready, reservation confirmed).',
      'The message and its status are recorded here.',
      'Staff can preview the exact text.',
      'Once a provider is connected, messages send for real.']),

    ('8.12  Staff', '/admin/staff', 'Managers, Owners',
     'Manages the people who work at the restaurant and their roles.',
     [('Staff list', 'Everyone with access, their role, and active status.'),
      ('Add / edit staff', 'Create accounts and assign roles (owner, manager, waiter, kitchen).'),
      ('Activate / deactivate', 'Turn access on or off for an employee.')],
     ['Add a new employee and set their role.',
      'That role decides which screens they can open.',
      'Deactivate anyone who leaves to revoke access.',
      'Keep the team list current and secure.']),

    ('8.13  Restaurants (Super-Admin)', '/admin/restaurants', 'Owners only',
     'The multi-restaurant control room - the heart of the SaaS model.',
     [('All tenants', 'Every restaurant in the system with its own stats.'),
      ('Switch active restaurant', 'Jump between restaurants you manage.'),
      ('Settings per restaurant', 'Edit tax rate, service charge, currency, and branding.')],
     ['The owner sees a list of all their restaurants.',
      'Each shows its own orders, revenue, and stats - fully separated.',
      'The owner switches the active restaurant with one click.',
      'Per-restaurant settings (tax, service charge, currency) are edited here.']),
]

for title_txt, route, who, purpose, features_list, steps in admin_screens:
    h2(title_txt)
    body('Route: %s  |  Who can use it: %s' % (route, who))
    h3('Purpose')
    body(purpose)
    h3('What is on the screen')
    for label, txt in features_list:
        bullet(txt, bold_lead=label + ': ')
    h3('How it is used - step by step')
    for s in steps:
        numbered(s)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 9. DATABASE
# ============================================================
h1('9.  Database & Backend (Supabase)')
body('All information is stored in a PostgreSQL database hosted on Supabase. The defining feature '
     'of the design is multi-tenancy: almost every table carries a restaurant_id column, so each '
     'restaurant\'s data is tagged and kept separate.')

h2('Main database tables')
make_table(['Table', 'What it stores'], [
    ('restaurants', 'Each tenant: name, slug, tagline, tax rate, service charge, currency.'),
    ('staff', 'Employees, their role, and the link to their login account.'),
    ('tables', 'Physical tables, seat counts, and live status.'),
    ('qr_codes', 'The unique QR token and link for each table.'),
    ('menu_categories', 'Menu sections and their display order.'),
    ('menu_items', 'Dishes: name, description, price, image, tags, availability.'),
    ('inventory_items', 'Stock: name, unit, quantity, low-stock threshold.'),
    ('menu_item_ingredients', 'Recipes linking dishes to ingredients for auto-deduct.'),
    ('orders / order_items', 'Orders and their line items, with totals and status.'),
    ('service_requests', 'Guest calls: waiter, water, bill, or assistance.'),
    ('reservations', 'Table bookings with guest details and status.'),
    ('customers', 'Customer contact records.'),
    ('bills', 'Final bills with the full money breakdown.'),
    ('feedback', 'Food / service / experience ratings and comments.'),
    ('upsell_rules', 'Which dish suggests which add-on, and the message.'),
    ('notifications', 'WhatsApp message log with recipient and status.'),
], widths=[2.2, 4.3])

h2('Security: Row Level Security (RLS)')
body('Database-level rules ensure staff can only ever read or change rows belonging to their own '
     'restaurant. Anonymous customers are allowed to read the public menu and tables, and to create '
     'orders, reservations, feedback, and service requests - but nothing else. This protection lives '
     'in the database itself, so it holds even if the user interface is bypassed.')

h2('The backend toggle')
body('If the Supabase keys are present, the app uses the live cloud database. If not, it '
     'automatically uses local browser storage seeded with two demo restaurants - "Spice Garden" '
     '(Indian) and "Cafe Aroma" (cafe) - each with staff, menus, tables, orders, and reservations. '
     'This is why the app works perfectly for demos out of the box.')

doc.add_page_break()

# ============================================================
# 10. LIFECYCLE
# ============================================================
h1('10.  Complete Order Lifecycle (End to End)')
body('This ties every dashboard together. Follow one order from the table to the owner\'s '
     'analytics:')

numbered('A guest scans the table QR code and the menu opens on their phone.', bold_lead='Customer: ')
numbered('They add dishes to the cart and place the order.', bold_lead='Customer: ')
numbered('The order instantly appears in the New column with a toast alert.', bold_lead='Kitchen board: ')
numbered('A cook taps Start (Preparing), then Ready when the food is done.', bold_lead='Kitchen board: ')
numbered('The dish shows under "Ready to serve"; a waiter carries it out and marks it served.', bold_lead='Waiter board: ')
numbered('The live timeline updates at each step - pending, preparing, ready, served.', bold_lead='Customer tracking: ')
numbered('Ingredients for the dishes are auto-deducted from stock.', bold_lead='Inventory: ')
numbered('A bill is generated with tax and service charge, ready for PDF export.', bold_lead='Billing: ')
numbered('The guest rates the meal; the rating lands in the Feedback screen.', bold_lead='Feedback: ')
numbered('The revenue and the dish sale flow into the owner\'s charts.', bold_lead='Analytics: ')

callout('THE BIG IDEA:', 'One customer tap ripples through the kitchen, the floor, the stockroom, '
        'the cash drawer, and the boardroom - automatically and in real time. That single connected '
        'flow is what SmartDine is all about.')

doc.add_page_break()

# ============================================================
# 11. HOW TO RUN
# ============================================================
h1('11.  How to Run the Project')
h2('Prerequisites')
bullet('Node.js installed (the JavaScript runtime).')
bullet('A code editor (e.g. VS Code).')
bullet('Optionally, a Supabase account to run against the live database.')

h2('Steps')
numbered('Open a terminal in the project folder.')
numbered('Install dependencies:  npm install', bold_lead='')
numbered('Start the development server:  npm run dev', bold_lead='')
numbered('Open the local address shown in the terminal in your browser.')
numbered('On the login page, tap any of the demo role buttons (Owner / Manager / Waiter / Kitchen) '
         'to explore that role.')
numbered('To go live, add VITE_SUPABASE_URL and VITE_SUPABASE_ANON_KEY to the environment file and '
         'run the SQL scripts in the supabase folder.')

h2('Useful commands')
make_table(['Command', 'What it does'], [
    ('npm run dev', 'Starts the development server.'),
    ('npm run build', 'Builds the production version of the app.'),
    ('npm run preview', 'Previews the production build locally.'),
    ('npm run lint', 'Checks the code for type errors.'),
], widths=[2.2, 4.3])

doc.add_page_break()

# ============================================================
# 12. GLOSSARY
# ============================================================
h1('12.  Glossary')
make_table(['Term', 'Meaning'], [
    ('SaaS', 'Software-as-a-Service - software delivered over the web to many customers.'),
    ('Multi-tenant', 'One system serving many separate customers (restaurants) with isolated data.'),
    ('Tenant', 'A single restaurant within the multi-tenant system.'),
    ('Dashboard', 'A screen that summarises information and lets a user act on it.'),
    ('KPI', 'Key Performance Indicator - a headline number like today\'s revenue.'),
    ('Kanban board', 'A board with columns that items move across (used by the kitchen).'),
    ('Toast', 'A small pop-up alert that appears briefly on screen.'),
    ('QR code', 'A scannable square code; here it links a phone to a specific table.'),
    ('Upsell', 'Suggesting an add-on to increase the order value.'),
    ('Realtime', 'Updates that appear instantly across all screens without refreshing.'),
    ('RLS', 'Row Level Security - database rules that restrict who can see which rows.'),
    ('Supabase', 'The cloud backend providing the database, auth, and realtime.'),
    ('Slug', 'A short text id for a restaurant used in the web address.'),
], widths=[1.6, 4.9])

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.paragraph_format.space_before = Pt(24)
r = end.add_run('- End of Document -')
r.bold = True
r.font.color.rgb = PRIMARY
r.font.size = Pt(12)

out = r'd:\SAMP IT\SmartDine\SmartDine_Project_Explanation.docx'
doc.save(out)
print('Saved:', out)
