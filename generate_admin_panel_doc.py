# -*- coding: utf-8 -*-
"""Generate the full SmartDine (ToDining) ADMIN PANEL document.

A standalone, deep-dive Word document covering only the Admin Panel:
the /admin-panel workspace manager, the /admin dashboard shell, role
visibility, and every one of the 13 admin screens, all matched to the
current source code.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Theme colors ----------
PRIMARY = RGBColor(0xC0, 0x39, 0x2B)   # restaurant red
DARK    = RGBColor(0x1F, 0x2A, 0x37)   # slate
ACCENT  = RGBColor(0xE6, 0x7E, 0x22)   # orange
GREY    = RGBColor(0x55, 0x5F, 0x6B)
LIGHT   = RGBColor(0xFA, 0xF3, 0xEF)

doc = Document()

# ---------- Base styles ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=10, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = PRIMARY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'E67E22')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = DARK
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT
    return p


def body(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(label, text, fill='FAF3EF', barcolor=PRIMARY):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + '  ')
    r.bold = True
    r.font.color.rgb = barcolor
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    doc.add_paragraph()
    return tbl


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], 'C0392B')
        set_cell_text(hdr[i], htext, bold=True, white=True, size=10)
    for r_i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
            if r_i % 2 == 1:
                shade(cells[i], 'F7F1ED')
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SmartDine')
r.bold = True
r.font.size = Pt(54)
r.font.color.rgb = PRIMARY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Admin Panel')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ACCENT

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run('The Management Back Office - Complete Guide')
r.font.size = Pt(15)
r.font.color.rgb = GREY

doc.add_paragraph()
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run('Workspace Manager  ' + chr(0x2022) + '  13 Admin Screens  ' + chr(0x2022) +
                 '  Step-by-Step')
r.font.size = Pt(12)
r.bold = True
r.font.color.rgb = DARK

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run('Dashboard  ' + chr(0x2022) + '  Analytics  ' + chr(0x2022) + '  Orders  ' + chr(0x2022) +
                 '  Tables & QR  ' + chr(0x2022) + '  Menu  ' + chr(0x2022) + '  Categories  ' + chr(0x2022) +
                 '  Reservations  ' + chr(0x2022) + '  Inventory  ' + chr(0x2022) + '  Billing  ' + chr(0x2022) +
                 '  Feedback  ' + chr(0x2022) + '  Notifications  ' + chr(0x2022) + '  Staff  ' + chr(0x2022) +
                 '  Restaurants')
r.font.size = Pt(10.5)
r.italic = True
r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Document Version 1.0\nPrepared: June 2026')
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
h1('Table of Contents')
toc_items = [
    '1.  What the Admin Panel Is',
    '2.  The Two Parts of the Admin Panel',
    '3.  Getting In - Access & Login Behaviour',
    '4.  The Workspace Manager (/admin-panel)',
    '5.  Creating a New Hotel Workspace - Step by Step',
    '6.  The Admin Dashboard Shell (/admin)',
    '7.  Navigation Map & Role Visibility',
    '8.  The 13 Admin Screens - One by One',
    '       8.1  Dashboard',
    '       8.2  Analytics (Owner)',
    '       8.3  Orders',
    '       8.4  Tables & QR',
    '       8.5  Menu',
    '       8.6  Categories',
    '       8.7  Reservations',
    '       8.8  Inventory',
    '       8.9  Billing',
    '       8.10 Feedback',
    '       8.11 Notifications',
    '       8.12 Staff',
    '       8.13 Restaurants (Owner)',
    '9.  Real-Time Sync - Why Edits Appear Everywhere',
    '10. Security Model',
    '11. Quick Reference',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    indent = item.startswith('     ')
    r = p.add_run(item.strip())
    r.font.size = Pt(11)
    if not indent and item[0].isdigit():
        r.bold = True
        r.font.color.rgb = DARK
    else:
        r.font.color.rgb = GREY
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)

doc.add_page_break()

# ============================================================
# 1. WHAT THE ADMIN PANEL IS
# ============================================================
h1('1.  What the Admin Panel Is')
body('The Admin Panel is the management back office of SmartDine - the part of the system that '
     'managers and owners use to run the business. While customers order from their phones and '
     'staff work the kitchen and waiter boards, the Admin Panel is where the restaurant is '
     'actually configured, monitored, and controlled.')

body('Everything the customer-facing app shows - the menu, the prices, the tables, the QR codes, '
     'the opening of a reservation slot - is set up and maintained here. And everything that '
     'happens out on the floor - orders, bills, feedback, stock movements - flows back into the '
     'Admin Panel so the manager always has the full picture.')

h2('Who it is for')
make_table(['Role', 'What they get in the Admin Panel'], [
    ('Owner', 'Everything - including the owner-only Analytics and Restaurants screens.'),
    ('Manager', 'Day-to-day running: all screens except the two owner-only ones.'),
    ('Waiter / Kitchen', 'No access - they are sent to their own boards instead.'),
], widths=[1.6, 4.9])

callout('IN SHORT:', 'The Admin Panel is the control room. If it can be configured, measured, or '
        'managed in SmartDine, it happens here.')

doc.add_page_break()

# ============================================================
# 2. TWO PARTS
# ============================================================
h1('2.  The Two Parts of the Admin Panel')
body('The Admin Panel is made of two distinct screens that work together. Understanding the split '
     'is the key to understanding the whole panel.')

h2('Part 1 - The Workspace Manager  (/admin-panel)')
body('A card-based landing page that lists every hotel (restaurant) as its own workspace. It does '
     'NOT open a restaurant directly. Instead it lets the admin pick which hotel to manage, or '
     'create a brand-new one. Each card shows that hotel\'s live stats (menu items, tables, '
     'orders, staff). Choosing a card and confirming enters that hotel\'s dedicated dashboard.')

h2('Part 2 - The Admin Dashboard  (/admin)')
body('The full management dashboard for one selected hotel. It has a sidebar with thirteen '
     'screens, a restaurant switcher at the top, and the main work area. Every screen here acts '
     'only on the currently active hotel - its data is completely isolated from the others.')

callout('THE FLOW:', 'Open /admin-panel -> pick or create a hotel workspace -> confirm -> land in '
        '/admin for that hotel. The workspace manager is the door; the dashboard is the room.')

h2('Why two screens instead of one')
body('SmartDine is multi-tenant: a single installation can run many restaurants at once. The '
     'workspace manager makes that explicit and safe. The admin first chooses exactly which hotel '
     'they intend to manage, so there is never any confusion about whose data is on screen. Once '
     'inside, the dashboard is focused entirely on that one hotel.')

doc.add_page_break()

# ============================================================
# 3. ACCESS & LOGIN
# ============================================================
h1('3.  Getting In - Access & Login Behaviour')
body('The Admin Panel is reachable in two ways, and its access rules are deliberately practical '
     'so demos and real use both work smoothly.')

h2('The two entry points')
make_table(['Path', 'What it opens'], [
    ('/admin-panel', 'The Workspace Manager - the card list of all hotels.'),
    ('/admin', 'The dashboard for the currently active hotel.'),
], widths=[1.8, 4.7])

h2('Open access')
body('The /admin route is "open": it can be reached straight from the workspace manager without '
     'logging in. When the panel is opened without a session, the visitor is treated as an owner, '
     'so the full menu - including the owner-only Analytics and Restaurants screens - is visible. '
     'This is what makes the panel instantly explorable.')

body('Importantly, "open" only relaxes the no-session case. A user who IS logged in but is not a '
     'manager or owner (for example a waiter) is still redirected away - they cannot reach the '
     'admin screens just by typing the address.')

h2('Signing in from inside the panel')
body('When no one is logged in, the sidebar shows a "Staff sign in" button instead of a user '
     'profile. Staff can authenticate at any time without leaving their place. Once signed in, the '
     'sidebar shows their name and role, and a log-out button. In production each staff member has '
     'a real account, and their role decides exactly which screens appear.')

callout('LANDING RULE:', 'After a real login, an Owner lands on Analytics and a Manager lands on '
        'Orders - each role starts on the screen most useful to their job.')

doc.add_page_break()

# ============================================================
# 4. WORKSPACE MANAGER
# ============================================================
h1('4.  The Workspace Manager  (/admin-panel)')
body('This is the front door of the Admin Panel. It presents every hotel as a workspace card and '
     'lets the admin manage which one to enter.')

h2('What is on the screen')
bullet('A header with the SmartDine wordmark and a "Back to website" link.', bold_lead='Header: ')
bullet('A title - "Your Hotel Workspaces" - explaining that each hotel is a fully isolated '
       'workspace with its own menu, orders, staff, tables and settings.', bold_lead='Intro: ')
bullet('One card per hotel, each showing the hotel\'s colour/logo, name, tagline, and four live '
       'stat tiles: Menu, Tables, Orders, Staff.', bold_lead='Workspace cards: ')
bullet('The card for the hotel you are currently in is marked "Active workspace"; the rest say '
       '"Open workspace".', bold_lead='Active marker: ')
bullet('A dashed "Create New Hotel Workspace" card to spin up a fresh, isolated hotel.',
       bold_lead='Create card: ')
bullet('A "Staff? Sign in here" link at the bottom for employees.', bold_lead='Staff link: ')

h2('The stats are live')
body('The four numbers on each card (menu items, tables, orders, staff) are read live for that '
     'specific hotel. Because the page re-reads whenever the data changes, a newly created '
     'workspace card - and any change to a hotel\'s counts - appears instantly without a refresh.')

h2('How an admin uses it - step by step')
numbered('Open /admin-panel. Every hotel you manage is shown as a card with its live stats.')
numbered('To manage an existing hotel, click its card. A confirmation dialog appears.')
numbered('Confirm. The app switches the active workspace to that hotel and opens its /admin dashboard.')
numbered('To add a hotel instead, click the dashed "Create New Hotel Workspace" card and fill the form.')

callout('WHY CONFIRM FIRST:', 'Switching workspace changes which hotel\'s data every admin screen '
        'shows. The confirmation step - "All data shown belongs only to this hotel" - makes the '
        'switch deliberate, preventing accidental edits to the wrong restaurant.')

doc.add_page_break()

# ============================================================
# 5. CREATE WORKSPACE
# ============================================================
h1('5.  Creating a New Hotel Workspace - Step by Step')
body('Adding a new restaurant to the system is done entirely from the workspace manager. It takes '
     'one short form.')

h2('The create form')
make_table(['Field', 'Purpose', 'Required'], [
    ('Hotel name', 'The restaurant\'s display name (e.g. "AK Restaurant").', 'Yes'),
    ('Owner email', 'The login identity for this hotel\'s owner.', 'Yes'),
    ('Username', 'An optional alternative login handle.', 'No'),
    ('Password', 'The owner\'s login password.', 'No'),
    ('Brand colour', 'The accent colour shown on the workspace card and logo tile.', 'No'),
    ('Description', 'A short note about the hotel, shown as its tagline.', 'No'),
], widths=[1.5, 3.8, 1.2])

h2('The steps')
numbered('On /admin-panel, click the dashed "Create New Hotel Workspace" card.')
numbered('Enter the hotel name (required) and the owner email or username (required for login).')
numbered('Optionally set a username, password, brand colour, and description.')
numbered('Click "Create workspace".')
numbered('The new hotel is created with default settings, categories and tables, and its card '
         'appears immediately in the list.')

callout('VALIDATION:', 'The form refuses to submit without a hotel name and an owner email/username '
        '- the two pieces needed to identify the hotel and let its owner log in. A clear on-screen '
        'message says which field is missing.')

doc.add_page_break()

# ============================================================
# 6. DASHBOARD SHELL
# ============================================================
h1('6.  The Admin Dashboard Shell  (/admin)')
body('Once a workspace is entered, every admin screen is rendered inside a shared frame called the '
     'dashboard layout. This frame is the same on every screen, so the navigation never moves.')

h2('Parts of the frame')
bullet('On desktop, a fixed left sidebar holds the wordmark, the restaurant switcher, the '
       'navigation menu, and the user chip at the bottom.', bold_lead='Sidebar: ')
bullet('A dropdown at the top of the sidebar to switch the active hotel without returning to the '
       'workspace manager.', bold_lead='Restaurant switcher: ')
bullet('The list of admin screens. The active screen is highlighted; owner-only items are hidden '
       'from managers.', bold_lead='Navigation menu: ')
bullet('At the bottom: the signed-in user\'s name and role with a log-out button, or a "Staff sign '
       'in" button when no one is logged in.', bold_lead='User chip: ')
bullet('On phones, the sidebar collapses into a hamburger menu that slides out as a drawer, so the '
       'full panel works on a phone.', bold_lead='Mobile drawer: ')
bullet('The selected screen renders in the large central work area.', bold_lead='Main area: ')

h2('The restaurant switcher')
body('The switcher lets an owner managing several hotels jump between them from inside the '
     'dashboard. Switching instantly re-points every screen at the newly selected hotel - the '
     'orders, menu, revenue and everything else update to that hotel\'s isolated data.')

doc.add_page_break()

# ============================================================
# 7. NAV MAP & ROLE VISIBILITY
# ============================================================
h1('7.  Navigation Map & Role Visibility')
body('The sidebar lists thirteen screens in a fixed order. Two of them are owner-only and are '
     'simply not shown to managers.')

make_table(['#', 'Screen', 'Route', 'Visible to'], [
    ('1', 'Dashboard', '/admin', 'Manager, Owner'),
    ('2', 'Analytics', '/admin/analytics', 'Owner only'),
    ('3', 'Orders', '/admin/orders', 'Manager, Owner'),
    ('4', 'Tables & QR', '/admin/tables', 'Manager, Owner'),
    ('5', 'Menu', '/admin/menu', 'Manager, Owner'),
    ('6', 'Categories', '/admin/categories', 'Manager, Owner'),
    ('7', 'Reservations', '/admin/reservations', 'Manager, Owner'),
    ('8', 'Inventory', '/admin/inventory', 'Manager, Owner'),
    ('9', 'Billing', '/admin/billing', 'Manager, Owner'),
    ('10', 'Feedback', '/admin/feedback', 'Manager, Owner'),
    ('11', 'Notifications', '/admin/notifications', 'Manager, Owner'),
    ('12', 'Staff', '/admin/staff', 'Manager, Owner'),
    ('13', 'Restaurants', '/admin/restaurants', 'Owner only'),
], widths=[0.4, 1.8, 2.3, 2.0])

callout('NOTE ON OPEN ACCESS:', 'When the panel is opened without logging in, the visitor is '
        'treated as an owner - so all thirteen screens, including Analytics and Restaurants, are '
        'shown. A logged-in manager sees eleven; a logged-in waiter or kitchen user sees none.')

doc.add_page_break()

# ============================================================
# 8. THE 13 SCREENS
# ============================================================
h1('8.  The 13 Admin Screens - One by One')
body('Below, every admin screen is explained the same way: its route and who can use it, its '
     'purpose, what is on the screen, and how it is used step by step.')

admin_screens = [
    ('8.1  Dashboard', '/admin', 'Managers, Owners',
     'The landing screen for managers - a quick health check of the whole hotel at a glance.',
     [('KPI cards', 'Headline numbers: total menu items, total orders, number of tables, and today\'s revenue.'),
      ('Recent orders', 'A live list of the most recent orders with their current status.'),
      ('Quick actions', 'Shortcut links to jump straight to menu, tables, orders, and other screens.')],
     ['Open the panel and immediately see how the hotel is doing today.',
      'Glance at the KPI cards for the big picture (revenue, order count).',
      'Scan recent orders to spot anything that needs attention.',
      'Use a quick-action link to jump to the screen you need.']),

    ('8.2  Analytics (Owner)', '/admin/analytics', 'Owners only',
     'The owner\'s decision-making screen - turns raw operations into business insight with charts.',
     [('Revenue KPIs', 'Today, this week, this month, and average order value.'),
      ('7-day revenue chart', 'An area chart showing the revenue trend across the last week.'),
      ('Peak hours chart', 'A bar chart revealing the busiest times of day.'),
      ('Top foods', 'A ranked list of the best-selling dishes.'),
      ('Reservations & ratings', 'A pie chart of reservation outcomes plus average ratings.')],
     ['The owner signs in and lands here automatically.',
      'Check revenue KPIs to see how the business is trending.',
      'Read the peak-hours chart to plan staffing.',
      'Use top-foods to decide menu promotions and pricing.']),

    ('8.3  Orders', '/admin/orders', 'Managers, Owners',
     'A complete, filterable record of every order with full status control.',
     [('All orders list', 'Every order across the hotel, newest first.'),
      ('Status management', 'Change or correct an order\'s status when needed.'),
      ('Order details', 'See items, quantities, totals, and the table for each order.')],
     ['Open Orders to see the full history and live queue.',
      'Filter or scan to find a specific order.',
      'Adjust an order\'s status if something needs correcting.',
      'Use it as the source of truth when a customer has a question.']),

    ('8.4  Tables & QR', '/admin/tables', 'Managers, Owners',
     'Manages the physical tables and generates the QR codes guests scan.',
     [('Live table status', 'Each table shown as available, reserved, or occupied in real time.'),
      ('QR generation', 'Create and view a printable QR code for every table.'),
      ('Table setup', 'Add tables and set how many seats each has.')],
     ['Add each physical table with its seat count.',
      'Generate the QR code for every table.',
      'Print and place the QR codes on the tables.',
      'Watch table status change live as guests are seated and served.']),

    ('8.5  Menu', '/admin/menu', 'Managers, Owners',
     'The full create/edit/delete control over every dish on the menu.',
     [('Menu item CRUD', 'Add, edit, and remove dishes.'),
      ('Pricing & details', 'Set price, description, image, and tags.'),
      ('Availability toggle', 'Mark a dish as available or sold out instantly.')],
     ['Add a new dish with its name, price, photo, and category.',
      'Edit details or pricing whenever they change.',
      'Toggle a dish off when it sells out - it disappears from the customer menu live.',
      'Remove dishes that are retired from the menu.']),

    ('8.6  Categories', '/admin/categories', 'Managers, Owners',
     'Organises the menu into sections (starters, mains, drinks, desserts...).',
     [('Category list', 'All menu sections in their display order.'),
      ('Add / edit / sort', 'Create sections and control the order they appear in.')],
     ['Create the categories the menu should be grouped into.',
      'Set their order so the menu reads logically.',
      'Assign dishes to categories from the Menu screen.']),

    ('8.7  Reservations', '/admin/reservations', 'Managers, Owners',
     'Receives and manages table bookings made through the public reservation form.',
     [('Booking list', 'All reservations with name, contact, date, time, and guest count.'),
      ('Confirm / reject', 'Approve or decline each booking.'),
      ('Status tracking', 'See which bookings are pending, confirmed, or cancelled.')],
     ['A guest submits the public reservation form.',
      'The booking appears here as "pending".',
      'The manager confirms or rejects it.',
      'The table status updates to "reserved" for the chosen time.']),

    ('8.8  Inventory', '/admin/inventory', 'Managers, Owners',
     'Tracks raw stock and automatically deducts ingredients as dishes are ordered.',
     [('Stock items', 'All ingredients with current quantity and unit.'),
      ('Low-stock alerts', 'Highlights items that have dropped below their threshold.'),
      ('Auto-deduct via recipes', 'When a dish is ordered, its recipe ingredients are subtracted automatically.')],
     ['Add each ingredient with its starting quantity and low-stock threshold.',
      'Link ingredients to dishes (recipes) so deductions are automatic.',
      'As orders come in, stock drops on its own.',
      'Restock when an item triggers a low-stock alert.']),

    ('8.9  Billing', '/admin/billing', 'Managers, Owners',
     'Holds the history of every bill and lets staff export them as PDF.',
     [('Bill history', 'Every completed bill with its full breakdown.'),
      ('Tax & service charge', 'Each bill shows subtotal, tax, service charge, and grand total.'),
      ('PDF export', 'Download or print any bill as a professional PDF.')],
     ['When an order completes, a bill is generated automatically.',
      'The bill includes subtotal, tax, and service charge.',
      'Find any past bill in the history.',
      'Export or reprint it as a PDF when needed.']),

    ('8.10  Feedback', '/admin/feedback', 'Managers, Owners',
     'Collects the ratings and comments customers leave after their meal.',
     [('Ratings list', 'Food, service, and experience scores (1-5) per visit.'),
      ('Comments', 'Free-text feedback from guests.'),
      ('Trends', 'A view of how satisfaction is trending over time.')],
     ['A customer rates their visit on the tracking page.',
      'The rating and comment appear here.',
      'Management reviews feedback to spot problems and praise.',
      'Recurring issues guide improvements to food and service.']),

    ('8.11  Notifications', '/admin/notifications', 'Managers, Owners',
     'A log and preview of the WhatsApp messages the system sends to customers.',
     [('Message log', 'Every notification with recipient, type, and status.'),
      ('Previews', 'See exactly what the customer message looks like.'),
      ('Channel ready', 'Built to connect to WhatsApp providers (Meta / Twilio).')],
     ['The system prepares a WhatsApp message (e.g. order ready, reservation confirmed).',
      'The message and its status are recorded here.',
      'Staff can preview the exact text.',
      'Once a provider is connected, messages send for real.']),

    ('8.12  Staff', '/admin/staff', 'Managers, Owners',
     'Manages the people who work at the hotel and their roles.',
     [('Staff list', 'Everyone with access, their role, and active status.'),
      ('Add / edit staff', 'Create accounts and assign roles (owner, manager, waiter, kitchen).'),
      ('Activate / deactivate', 'Turn access on or off for an employee.')],
     ['Add a new employee and set their role.',
      'That role decides which screens they can open.',
      'Deactivate anyone who leaves to revoke access.',
      'Keep the team list current and secure.']),

    ('8.13  Restaurants (Owner)', '/admin/restaurants', 'Owners only',
     'The multi-restaurant control room - the heart of the SaaS model.',
     [('All tenants', 'Every hotel in the system with its own stats.'),
      ('Switch active hotel', 'Jump between the hotels you manage.'),
      ('Settings per hotel', 'Edit tax rate, service charge, currency, and branding.')],
     ['The owner sees a list of all their hotels.',
      'Each shows its own orders, revenue, and stats - fully separated.',
      'The owner switches the active hotel with one click.',
      'Per-hotel settings (tax, service charge, currency) are edited here.']),
]

for title_txt, route, who, purpose, features_list, steps in admin_screens:
    h2(title_txt)
    body('Route: %s  |  Who can use it: %s' % (route, who))
    h3('Purpose')
    body(purpose)
    h3('What is on the screen')
    for label, txt in features_list:
        bullet(txt, bold_lead=label + ': ')
    h3('How it is used - step by step')
    for s in steps:
        numbered(s)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 9. REAL-TIME SYNC
# ============================================================
h1('9.  Real-Time Sync - Why Edits Appear Everywhere')
body('A defining behaviour of the Admin Panel is that an edit made here reaches every other screen '
     'instantly - the live customer menu, the kitchen and waiter boards, and the other admin '
     'screens - without anyone refreshing.')

h2('How it works')
numbered('Every admin screen reads and writes through a shared set of services (menu service, '
         'order service, and so on) - never the database directly.', bold_lead='Shared services: ')
numbered('Whenever a service changes data, it emits a change event on an internal event bus.',
         bold_lead='Change events: ')
numbered('Screens subscribe to those events with a live-query hook, so they automatically re-read '
         'and re-render the moment something they care about changes.', bold_lead='Live queries: ')

body('This is exactly why the workspace manager\'s stat cards update the instant a workspace is '
     'created, and why toggling a dish to "sold out" in the Menu screen removes it from the '
     'customer\'s phone immediately.')

callout('THE CONTRACT:', 'Admin edits reach the live customer site because services emit change '
        'events and screens listen for them. Any new admin action that changes data must go '
        'through a service so the rest of the app stays in sync.')

doc.add_page_break()

# ============================================================
# 10. SECURITY MODEL
# ============================================================
h1('10.  Security Model')
body('The Admin Panel is protected at two independent levels, so a gap in one is still covered by '
     'the other.')

h2('Level 1 - The Role Guard (in the screen)')
body('Every admin route is wrapped in a Role Guard that only lets managers and owners through. '
     'The /admin route is marked "open", which permits an anonymous visitor in for demo and quick '
     'access - but a logged-in user with a non-admin role (waiter, kitchen) is still redirected '
     'away. Within the panel, owner-only screens (Analytics, Restaurants) are hidden from managers.')

h2('Level 2 - Row Level Security (in the database)')
body('In production, the database enforces that staff can only ever read or change rows belonging '
     'to their own hotel. Because this lives in the database itself, it holds even if the screen '
     'protection were somehow bypassed. This is what guarantees one hotel can never see another\'s '
     'data.')

h2('Tenant isolation')
body('Almost every record carries a hotel (restaurant) id. The active workspace - chosen in the '
     'workspace manager or the restaurant switcher - decides which hotel\'s data every admin screen '
     'reads and writes. That is how a single installation runs many restaurants while keeping each '
     'one\'s data completely private.')

callout('TWO LOCKS, ONE DOOR:', 'The Role Guard controls who reaches the screens; Row Level '
        'Security controls what data those screens can ever touch. Together they make cross-hotel '
        'access effectively impossible.')

doc.add_page_break()

# ============================================================
# 11. QUICK REFERENCE
# ============================================================
h1('11.  Quick Reference')

h2('Routes at a glance')
make_table(['Route', 'Screen', 'Access'], [
    ('/admin-panel', 'Workspace Manager (card list)', 'Open'),
    ('/admin', 'Dashboard', 'Manager, Owner'),
    ('/admin/analytics', 'Analytics', 'Owner'),
    ('/admin/orders', 'Orders', 'Manager, Owner'),
    ('/admin/tables', 'Tables & QR', 'Manager, Owner'),
    ('/admin/menu', 'Menu', 'Manager, Owner'),
    ('/admin/categories', 'Categories', 'Manager, Owner'),
    ('/admin/reservations', 'Reservations', 'Manager, Owner'),
    ('/admin/inventory', 'Inventory', 'Manager, Owner'),
    ('/admin/billing', 'Billing', 'Manager, Owner'),
    ('/admin/feedback', 'Feedback', 'Manager, Owner'),
    ('/admin/notifications', 'Notifications', 'Manager, Owner'),
    ('/admin/staff', 'Staff', 'Manager, Owner'),
    ('/admin/restaurants', 'Restaurants', 'Owner'),
], widths=[2.2, 2.5, 1.8])

h2('Common admin tasks - where to do them')
make_table(['I want to...', 'Go to'], [
    ('Add or switch a hotel', 'Workspace Manager (/admin-panel) or the restaurant switcher'),
    ('Add a dish or mark it sold out', 'Menu'),
    ('Group dishes into sections', 'Categories'),
    ('Add a table and print its QR', 'Tables & QR'),
    ('Confirm or reject a booking', 'Reservations'),
    ('Check today\'s revenue and trends', 'Dashboard / Analytics'),
    ('Correct an order\'s status', 'Orders'),
    ('Reprint or export a bill', 'Billing'),
    ('Track stock and restock', 'Inventory'),
    ('Read customer ratings', 'Feedback'),
    ('Add or deactivate an employee', 'Staff'),
    ('Edit tax, service charge, currency', 'Restaurants'),
], widths=[3.0, 3.5])

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.paragraph_format.space_before = Pt(24)
r = end.add_run('- End of Document -')
r.bold = True
r.font.color.rgb = PRIMARY
r.font.size = Pt(12)

out = r'd:\SAMP IT\SmartDine\SmartDine_Admin_Panel.docx'
doc.save(out)
print('Saved:', out)
